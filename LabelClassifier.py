from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler
from sklearn.svm import SVC
from sklearn.feature_extraction.text import TfidfVectorizer
import pandas as pd
import os
from openpyxl import load_workbook
from docx import Document

# Load data from spreadsheet
data = pd.read_excel('C:\\Users\\adij2\\Downloads\\question_label_variations_expanded.xlsx')

# Assign labels
X = data['raw_label'].astype(str)         # e.g., "Q1", "question one"
y = data['canonical_label'].astype(str)   # ensure labels are strings

# Train-test split
X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, random_state=42
)

# Text vectorization with TF-IDF
vectorizer = TfidfVectorizer(analyzer='char_wb', ngram_range=(2, 4))
X_train_tfidf = vectorizer.fit_transform(X_train)
X_test_tfidf = vectorizer.transform(X_test)

# Train SVM
model = SVC(kernel='rbf', C=1.0, gamma='scale')
model.fit(X_train_tfidf, y_train)

# ====== Load raw labels from Word ======
raw_doc = Document("C:\\Users\\adij2\\Downloads\\raw_labels.docx")
raw_labels = [p.text.strip() for p in raw_doc.paragraphs if p.text.strip()]
raw_df = pd.DataFrame({"raw_label": raw_labels})

# Predict canonical
raw_df["predicted_canonical"] = model.predict(vectorizer.transform(raw_df["raw_label"]))

# ====== Save ProcessedLabels ======
processed_doc = Document()
table = processed_doc.add_table(rows=1, cols=2)
hdr = table.rows[0].cells
hdr[0].text = "Raw Label"
hdr[1].text = "Predicted Canonical Label"

for _, row in raw_df.iterrows():
    cells = table.add_row().cells
    cells[0].text = str(row["raw_label"])
    cells[1].text = str(row["predicted_canonical"])

processed_doc.save("C:\\Users\\adij2\\Downloads\\ProcessedLabels.docx")
print("Processed Labels saved")

gt_doc_path = "C:\\Users\\adij2\\Downloads\\GroundTruth.docx"
raw_labels_set = set(raw_df["raw_label"].astype(str))

# Load existing GroundTruth doc
gt_doc = Document(gt_doc_path)

# If a table exists, use it; otherwise create a new table with headers
if gt_doc.tables:
    table = gt_doc.tables[0]
else:
    table = gt_doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Ground Truth"
    table.rows[0].cells[1].text = "Found/Not Found"

# Collect existing ground truth labels in the table
existing_labels = set()
for row in table.rows[1:]:  # skip header
    gt_label = row.cells[0].text.strip()
    existing_labels.add(gt_label)

# Gather all labels to process: existing + any new ones from elsewhere
# Here, just re-using existing table cells (or could append new ones)
for row in table.rows[1:]:
    gt_label = row.cells[0].text.strip()
    row.cells[1].text = "Found" if gt_label in raw_labels_set else "Not Found"

gt_doc.save(gt_doc_path)
print("Ground Truth updated with Found/Not Found")
